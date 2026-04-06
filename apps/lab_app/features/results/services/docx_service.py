# apps/lab_app/features/results/services/docx_service.py
from __future__ import annotations

from datetime import datetime, date
from pathlib import Path
from typing import Any

import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from apps.lab_app.features.results.services.compute_service import compute_grid_flags_and_apply


def _add_logo_or_raise(doc: Document, logo_path: str) -> None:
    logo_path = (logo_path or "").strip()
    if not logo_path:
        raise RuntimeError("logo_path is empty in lab_profile")

    logo_path = os.path.abspath(logo_path)
    if not os.path.exists(logo_path):
        raise RuntimeError(f"logo_path does not exist: {logo_path}")

    section = doc.sections[0]
    header = section.header

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.1))


def _add_center_lines(doc: Document, lines: list[str]) -> None:
    for ln in lines:
        s = str(ln or "").strip()
        if not s:
            continue
        p = doc.add_paragraph(s)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_patient_block(doc: Document, patient_row: dict, printed_dt: str) -> None:
    pid = patient_row.get("Patient ID", "-")
    name = patient_row.get("Name", "-")
    sex = patient_row.get("Sex", "-")
    age = patient_row.get("Age", "-")

    doc.add_paragraph("")  # spacer

    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"

    t.cell(0, 0).text = f"Name: {name}"
    t.cell(0, 1).text = f"Printed: {printed_dt}"

    t.cell(1, 0).text = f"Patient ID: {pid}"
    t.cell(1, 1).text = f"Sex: {sex}"

    t.cell(2, 0).text = f"Age: {age}"
    t.cell(2, 1).text = ""


def _bold_first_paragraph(cell) -> None:
    # Safely bold whatever exists in the first paragraph of a cell
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    if not p.runs:
        # create a run so we can bold it without crashing
        p.add_run("")
    for r in p.runs:
        r.font.bold = True


def _render_structured(doc: Document, payload: dict) -> None:
    rows = payload.get("rows", []) or []

    # Always include flag column
    data = [["Parameter", "Result", "Unit", "Ref. Range", "Flag"]]
    for r in rows:
        data.append([
            str(r.get("parameter", "") or ""),
            str(r.get("result", "") or ""),
            str(r.get("unit", "") or ""),
            str(r.get("ref_range", "") or ""),
            str(r.get("flag", "") or ""),
        ])

    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"

    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.cell(i, j).text = str(val)

    # bold header
    for cell in table.rows[0].cells:
        _bold_first_paragraph(cell)


def _render_grid(doc: Document, payload: dict) -> None:
    # Prefer UIX snapshot + values
    uix = payload.get("uix") or {}
    snap = (uix.get("template_snapshot") or payload.get("grid") or {}) or {}
    values = (uix.get("values") or {}) or {}
    grid = payload.get("grid") or {}

    cells = values.get("cells") or grid.get("cells") or []

    if not isinstance(cells, list) or not cells:
        doc.add_paragraph("(No table data)")
        return

    # Apply schema-based flag recompute (writes back into values["cells"])
    try:
        values2, _flags = compute_grid_flags_and_apply(
            template_snapshot=snap,
            values={"cells": cells},
        )
        cells = values2.get("cells") or cells
    except Exception:
        # never break export due to flag recompute
        pass

    ncols = max((len(r) for r in cells if isinstance(r, list)), default=0)
    if ncols <= 0:
        doc.add_paragraph("(No table data)")
        return

    # Pad ragged rows safely; render exactly saved cells, no numbering injection
    padded: list[list[str]] = []
    for r in cells:
        r = r if isinstance(r, list) else []
        padded.append([str(x or "") for x in (r + [""] * (ncols - len(r)))])

    table = doc.add_table(rows=len(padded), cols=ncols)
    table.style = "Table Grid"

    for i, row in enumerate(padded):
        for j, val in enumerate(row):
            table.cell(i, j).text = str(val)

    # Make row 0 bold like PDF header
    if table.rows:
        for cell in table.rows[0].cells:
            _bold_first_paragraph(cell)


def generate_bundle_docx(
    output_path: str,
    lab_profile: dict,
    patient_row: dict,
    bundle_results: dict[str, dict],
) -> str:
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Single logo only (optional)
    try:
        _add_logo_or_raise(doc, str(lab_profile.get("logo_path") or ""))
    except Exception as e:
        # Keep it non-fatal
        print("DOCX logo skipped:", e)

    lab_name = lab_profile.get("lab_name", "Laboratory")
    address = lab_profile.get("address", "")
    phone = lab_profile.get("phone", "")
    email = lab_profile.get("email", "")

    header_line = "  ".join([x for x in [str(phone).strip(), str(email).strip()] if x])
    _add_center_lines(doc, [str(lab_name), str(address), header_line])

    doc.add_paragraph("")  # spacer
    title = doc.add_paragraph("Patient Report (Bundle)")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.bold = True
        run.font.size = Pt(12)

    printed_dt = datetime.now().strftime("%Y-%m-%d %H:%M")
    _add_patient_block(doc, patient_row, printed_dt)

    doc.add_paragraph("")  # spacer

    for rid, payload in bundle_results.items():
        test_name = (payload.get("request", {}) or {}).get("test_name", "Test")

        h = doc.add_paragraph(f"{test_name} ({rid})")
        for run in h.runs:
            run.font.bold = True

        typ = (payload.get("type") or "").strip().lower()

        if typ == "structured":
            _render_structured(doc, payload)
        elif typ == "table":
            _render_grid(doc, payload)
        elif typ == "written":
            text = payload.get("text", {}) or {}
            for label, key in [
                ("Clinical Findings", "findings"),
                ("Interpretation", "interpretation"),
                ("Impression / Conclusion", "impression"),
                ("Recommendations", "recommendations"),
            ]:
                body = (text.get(key) or "").strip()
                if not body:
                    continue
                p = doc.add_paragraph(f"{label}: ")
                if p.runs:
                    p.runs[0].bold = True
                doc.add_paragraph(body)
        elif typ == "pc_template":
            rendered = (payload.get("rendered") or "").strip()
            doc.add_paragraph(rendered if rendered else "(No template content)")
        else:
            doc.add_paragraph("(Unsupported result type)")

        doc.add_paragraph("")  # spacer between tests

    # Signature footer
    scientist = (lab_profile.get("scientist_name") or "").strip()
    qual = (lab_profile.get("scientist_qualification") or "").strip()
    notes = (lab_profile.get("report_notes") or "").strip()

    doc.add_paragraph("")
    doc.add_paragraph("__________________________")
    if scientist:
        doc.add_paragraph(scientist)
    if qual:
        doc.add_paragraph(qual)
    if notes:
        doc.add_paragraph(f"Notes: {notes}")

    doc.add_paragraph("")
    doc.add_paragraph("Authorized Laboratory Report — Generated by Solunex Technologies Software")

    doc.save(str(out))
    return str(out)
